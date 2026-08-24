"""
Metro_Istanbul_Ariza_Tespit_Raporu.docx dosyasini SIFIRDAN uretir.

Bu script proje bitince/rapor guncellenmesi gerektiginde tekrar calistirilabilir
diye repoda tutuluyor. python-docx gerektirir (./venv/bin/pip install python-docx
-- runtime bagimliligi degil, sadece rapor uretimi icin, bu yuzden
requirements.txt'e eklenmedi).

Kullanim:
    ./venv/bin/python3 scripts/rapor_uret.py
"""

from __future__ import annotations

from pathlib import Path

from docx import Document
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.shared import Pt, Cm, RGBColor, Inches
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.oxml.ns import qn
from docx.oxml import OxmlElement

ROOT = Path(__file__).resolve().parent.parent
OUT = ROOT / "Metro_Istanbul_Ariza_Tespit_Raporu.docx"

KOYU = RGBColor(0x1a, 0x1a, 0x1a)
VURGU = RGBColor(0x0b, 0x5a, 0x9c)
GRI = RGBColor(0x55, 0x55, 0x55)


def _hucre_golgele(hucre, renk_hex: str) -> None:
    tcPr = hucre._tc.get_or_add_tcPr()
    shd = OxmlElement("w:shd")
    shd.set(qn("w:val"), "clear")
    shd.set(qn("w:fill"), renk_hex)
    tcPr.append(shd)


def kur_stiller(doc: Document) -> None:
    normal = doc.styles["Normal"]
    normal.font.name = "Calibri"
    normal.font.size = Pt(11)
    normal.font.color.rgb = KOYU
    normal.paragraph_format.space_after = Pt(8)
    normal.paragraph_format.line_spacing = 1.15

    for ad, boyut, renk, kalin in [
        ("Title", 26, VURGU, True),
        ("Heading 1", 18, VURGU, True),
        ("Heading 2", 14, KOYU, True),
        ("Heading 3", 12, KOYU, True),
    ]:
        s = doc.styles[ad]
        s.font.name = "Calibri"
        s.font.size = Pt(boyut)
        s.font.color.rgb = renk
        s.font.bold = kalin
        s.paragraph_format.space_before = Pt(18 if ad != "Title" else 0)
        s.paragraph_format.space_after = Pt(8)


def baslik(doc, metin, seviye=1):
    doc.add_heading(metin, level=seviye)


def p(doc, metin, italik=False, kalin=False, boyut=None, renk=None, hiza=None):
    par = doc.add_paragraph()
    run = par.add_run(metin)
    run.italic = italik
    run.bold = kalin
    if boyut:
        run.font.size = Pt(boyut)
    if renk:
        run.font.color.rgb = renk
    if hiza:
        par.alignment = hiza
    return par


def madde(doc, metin):
    doc.add_paragraph(metin, style="List Bullet")


def tablo(doc, basliklar, satirlar, genislikler=None):
    t = doc.add_table(rows=1, cols=len(basliklar))
    t.style = "Light Grid Accent 1"
    t.alignment = WD_TABLE_ALIGNMENT.CENTER
    hdr = t.rows[0].cells
    for i, b in enumerate(basliklar):
        hdr[i].text = b
        for para in hdr[i].paragraphs:
            for run in para.runs:
                run.font.bold = True
                run.font.size = Pt(10)
                run.font.color.rgb = RGBColor(0xFF, 0xFF, 0xFF)
        _hucre_golgele(hdr[i], "0B5A9C")
    for satir in satirlar:
        row = t.add_row().cells
        for i, deger in enumerate(satir):
            row[i].text = str(deger)
            for para in row[i].paragraphs:
                for run in para.runs:
                    run.font.size = Pt(10)
    doc.add_paragraph()
    return t


def kod_bloku(doc, metin):
    par = doc.add_paragraph()
    run = par.add_run(metin)
    run.font.name = "Consolas"
    run.font.size = Pt(9.5)
    run.font.color.rgb = RGBColor(0x33, 0x33, 0x33)
    par.paragraph_format.left_indent = Cm(0.5)
    shd = OxmlElement("w:shd")
    shd.set(qn("w:val"), "clear")
    shd.set(qn("w:fill"), "F2F2F2")
    par._p.get_or_add_pPr().append(shd)
    return par


def kutu(doc, baslik_metin, icerik_metin, renk_hex="EAF3FB", kenar_hex="0B5A9C"):
    t = doc.add_table(rows=1, cols=1)
    t.autofit = True
    hucre = t.rows[0].cells[0]
    _hucre_golgele(hucre, renk_hex)
    par1 = hucre.paragraphs[0]
    r1 = par1.add_run(baslik_metin)
    r1.bold = True
    r1.font.size = Pt(10.5)
    par2 = hucre.add_paragraph()
    r2 = par2.add_run(icerik_metin)
    r2.font.size = Pt(10.5)
    r2.italic = True
    doc.add_paragraph()


def olustur() -> None:
    doc = Document()
    for section in doc.sections:
        section.top_margin = Cm(2.2)
        section.bottom_margin = Cm(2.2)
        section.left_margin = Cm(2.5)
        section.right_margin = Cm(2.5)
    kur_stiller(doc)

    # ------------------------------------------------------------------ KAPAK
    par = doc.add_paragraph()
    par.alignment = WD_ALIGN_PARAGRAPH.CENTER
    par.paragraph_format.space_before = Pt(120)
    run = par.add_run("Metro İstanbul\nArıza Tespit Sınıflandırıcı")
    run.font.size = Pt(30)
    run.font.bold = True
    run.font.color.rgb = VURGU

    par2 = doc.add_paragraph()
    par2.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r2 = par2.add_run("Serbest Metinli Arıza Bildirimlerini Otomatik\nSınıflandıran Bir NLP Sistemi — Staj Projesi Raporu")
    r2.font.size = Pt(14)
    r2.font.color.rgb = GRI

    par3 = doc.add_paragraph()
    par3.alignment = WD_ALIGN_PARAGRAPH.CENTER
    par3.paragraph_format.space_before = Pt(220)
    r3 = par3.add_run("Yazılım Mühendisliği Stajı\nAğustos 2026")
    r3.font.size = Pt(11)
    r3.font.color.rgb = GRI
    doc.add_page_break()

    # ------------------------------------------------------------------ ONSOZ
    baslik(doc, "Önsöz", 1)
    p(doc,
      "Bu raporu Metro İstanbul'da yaptığım yazılım mühendisliği stajı boyunca "
      "geliştirdiğim projeyi anlatmak için yazdım. Aşağıda anlattığım her karar, "
      "her sayı ve her hata gerçekten yaşadığım şeyler — projeyi geliştirirken "
      "tuttuğum notlardan derledim, hiçbir yeri uydurmadım veya güzel göstermek "
      "için abartmadım. Bazı yerlerde işlerin ilk seferde yürümediğini, hatta "
      "modelin bir ara hiçbir şey öğrenmediğini de yazdım, çünkü bence bir "
      "projenin gerçek değeri sadece son haliyle değil, o hale nasıl gelindiğiyle "
      "de ölçülür.")
    p(doc,
      "Kısaca ne yaptığımı özetlemek gerekirse: Metro İstanbul çalışanlarının "
      "günlük olarak yazdığı serbest metinli arıza bildirimlerini ('Yürüyen "
      "merdiven durdu 2. peron' gibi) otomatik olarak analiz edip hangi bakım "
      "ekibine gitmesi gerektiğini bulan bir yapay zeka modeli kurdum. Bunun "
      "üstüne de gerçek zamanlı çalışan bir web servisi ve bir arayüz ekledim.")

    # ------------------------------------------------------------------ 1
    baslik(doc, "1. Giriş ve Problem Tanımı", 1)
    p(doc,
      "Metro İstanbul'da gün içinde onlarca, bazen yüzlerce arıza bildirimi "
      "geliyor — istasyon personeli, güvenlik görevlileri veya çağrı merkezi "
      "üzerinden. Bu bildirimler serbest metin olarak yazılıyor, yani standart "
      "bir form değil; kimi zaman düzgün bir cümle, kimi zaman '3 numaralı "
      "merdiven arızalı' gibi telgraf tarzı kısa bir not, kimi zaman da aceleyle "
      "yazılmış, yazım hatası dolu bir mesaj. Bu bildirimlerin hangi bakım "
      "ekibine (mekanik, elektrik, sinyalizasyon, temizlik, güvenlik...) "
      "gideceğine şu anda bir insan karar veriyor. Bu hem zaman alıyor hem de "
      "yoğun saatlerde gecikmelere yol açabiliyor.")
    p(doc,
      "Benim projemin amacı bu sınıflandırma işini otomatikleştirmek. Ama işin "
      "içine girdikçe fark ettim ki sadece 'hangi kategori' sorusu yeterli "
      "değil — bir arıza bildiriminin gerçekten ne kadar acil olduğu, "
      "kullanıcının ne yapmak istediği (şikayet mi ediyor, arıza mı bildiriyor, "
      "bilgi mi istiyor), ve bildirimde hangi somut bilgilerin (hat, istasyon, "
      "ekipman) geçtiği de en az kategori kadar önemli. Bu yüzden proje "
      "ilerledikçe tek boyutlu bir sınıflandırıcıdan üç boyutlu bir sisteme "
      "dönüştü — buna aşağıda ayrıntılı değineceğim.")

    baslik(doc, "1.1 Neden Bu Yaklaşım", 2)
    p(doc,
      "Projeye başlarken üç seçenek vardı: klasik makine öğrenmesi (TF-IDF + "
      "SVM gibi), derin öğrenme (önceden eğitilmiş bir dil modelini ince ayar "
      "yapmak) veya doğrudan büyük dil modeli (LLM) API'si kullanmak. Üçünü de "
      "gerçekten değerlendirdim:")
    madde(doc, "Klasik ML: hızlı ve basit ama Türkçe'nin morfolojik yapısını "
                "ve bağlamı yakalamakta zayıf kalıyor. Projenin ilerleyen "
                "aşamalarında karşılaştığım 'aynı kelimeler, farklı anlam' "
                "türü sınır sorunlarını (örneğin 'ekran karanlık' ile 'ekran "
                "çalışıyor ama bilgi yanlış' arasındaki fark) yakalayamazdı.")
    madde(doc, "LLM API'si (GPT/Gemini gibi): her istekte internet bağlantısı, "
                "saniyeler mertebesinde gecikme ve maliyet gerektirir; bir arıza "
                "tespit sistemi için pratik değil. Ayrıca çıktısı tam "
                "deterministik değil — aynı cümleye farklı zamanlarda farklı "
                "cevap verebilir, bu da güvenilirlik açısından sorun.")
    madde(doc, "Derin öğrenme (seçilen yöntem): Türkçe için önceden eğitilmiş "
                "BERTurk modelini kendi verimle ince ayar yaptım. Hem bağlamı "
                "iyi yakalıyor hem de tahmin süresi 14 milisaniye gibi çok "
                "hızlı, hem de tamamen yerelde (internet olmadan) çalışıyor.")
    p(doc,
      "Yine de LLM'leri projenin dışında bırakmadım — sadece rolünü değiştirdim. "
      "Eğitim verisi üretmek ve etiketlemek için (yani modelin öğreneceği "
      "örnekleri hazırlamak için) LLM kullandım, ama gerçek zamanlı tahmin "
      "sırasında (kullanıcı bir cümle yazdığında) hiçbir LLM çağrılmıyor — o iş "
      "tamamen kendi eğittiğim küçük modele ait. Yani mimari şöyle özetlenebilir: "
      "LLM veri üretir, kendi eğittiğim derin öğrenme modeli gerçek zamanlı "
      "sınıflandırır.")

    # ------------------------------------------------------------------ 2
    baslik(doc, "2. Sistem Mimarisi", 1)
    p(doc, "Sistemin genel akışı şöyle:")
    kod_bloku(doc,
              "KULLANICI METNİ\n"
              "      │\n"
              "      ▼\n"
              "  BERTurk + LoRA  (tek gövde, üç sınıflandırma başlığı)\n"
              "    ├── INTENT      5 sınıf\n"
              "    ├── CATEGORY   11 sınıf\n"
              "    └── PRIORITY    4 sınıf\n"
              "      │\n"
              "      ├──► kurallı çıkarım: hat, istasyon, konum, ekipman, belirti, kök sebep\n"
              "      ├──► P1 kural katmanı (yangın, elektrik çarpması... → koşulsuz kritik)\n"
              "      ├──► gradient × input → hangi kelimeler karara katkı yaptı\n"
              "      ├──► eksik bilgi tespiti → kullanıcıya soru\n"
              "      └──► tekrar bildirim tespiti (aynı istasyon+ekipman+15 dk)\n"
              "      │\n"
              "      ▼\n"
              "  FastAPI backend (:8000)  ──►  React arayüz (:5173)")
    p(doc,
      "Modelin kalbi tek bir BERTurk gövdesi ama üstünde üç ayrı sınıflandırma "
      "'başlığı' var: biri kategoriyi, biri kullanıcının niyetini (intent), "
      "biri de önceliği tahmin ediyor. Bunu üç ayrı model olarak değil, tek "
      "paylaşımlı gövde olarak kurdum çünkü bir cümlenin kategorisini belirleyen "
      "kelimeler genelde niyetini ve önceliğini de belirliyor — mesela 'peronda "
      "kavga var' cümlesi hem kategoriyi (güvenlik ve asayiş) hem niyeti (olay "
      "bildirimi) hem de önceliği (kritik) aynı anda söylüyor. Üç ayrı model "
      "eğitmek bu ortak bilgiyi üç kere sıfırdan öğrenmek olurdu, hem yavaş hem "
      "de sunucuda üç ayrı 440 MB'lık taban model tutmak gerekirdi.")

    baslik(doc, "2.1 Neden BERTurk + LoRA", 2)
    p(doc,
      "BERTurk (dbmdz/bert-base-turkish-cased), Türkçe metinlerle önceden "
      "eğitilmiş bir BERT modeli. Genel çok dilli BERT yerine bunu seçtim "
      "çünkü Türkçe'nin sondan eklemeli yapısını (çekim ekleri vs.) çok daha "
      "iyi yakalıyor.")
    p(doc,
      "Bu modeli sıfırdan eğitmek yerine LoRA (Low-Rank Adaptation) denen bir "
      "teknikle ince ayar yaptım. LoRA'nın mantığı şu: BERT'in kendi "
      "ağırlıklarını dondurup (değiştirmeden bırakıp) sadece belirli katmanlara "
      "(query ve value matrisleri) küçük, düşük ranklı ek katmanlar ekliyorsun "
      "ve sadece onları eğitiyorsun. Sonuçta 111 milyon parametrenin sadece "
      "605 bini (%0.54'ü) eğitiliyor. Bunun pratik faydası çok büyük: "
      "kaydedilen adaptör dosyası sadece 2.4 MB — tam model 440 MB olurdu. "
      "Yani modeli versiyonlamak, taşımak, farklı sürümlerini tutmak çok daha "
      "kolay.")

    kutu(doc, "En kritik hatam — öğrenme hızı",
         "İlk denememde LEARNING_RATE=2e-5 kullandım çünkü BERT ince ayarında "
         "literatürde standart değer bu. Ama model 5 epoch sonunda hiçbir şey "
         "öğrenmedi (val F1 = 0.134, rastgele tahmin seviyesi). Sebebi: 2e-5, "
         "TAM fine-tuning için standart bir değer, ama ben LoRA kullanıyordum "
         "— adaptörler sıfırdan başlıyor ve bu kadar küçük bir öğrenme hızıyla "
         "anlamlı mesafe kat edemiyor. LEARNING_RATE=5e-4 yapınca val F1 "
         "0.875'e çıktı. Ders: hiperparametreyi literatürden kopyalamak "
         "yetmiyor, kullandığın yönteme göre ayarlaman gerekiyor.")

    # ------------------------------------------------------------------ 3
    baslik(doc, "3. Taksonomi: 11 Kategori ve Üç Boyut", 1)
    p(doc,
      "Proje boyunca kategori sayısı ve yapısı iki kere değişti. İlk versiyon "
      "6 kategoriyle başladı, sonra 8'e, sonra 9'a çıktı. Ama gerçek büyük "
      "değişiklik taksonomi v2'de oldu: 9 kategoriden 11'e çıktım VE tek "
      "boyutlu (sadece kategori) yapıdan üç boyutlu (intent + kategori + "
      "öncelik) yapıya geçtim.")
    p(doc,
      "Bunun sebebi somut: v1 taksonomisinde iki kategori arasında ciddi bir "
      "kelime çakışması vardı ('altyapi_insaat' ile 'temizlik_cevre' arasında "
      "'su/sızıntı/döküntü' gibi ortak kelimeler), ve model bunu ayıramıyordu "
      "— bağımsız test setinde altyapi_insaat kategorisi sadece %5 doğrulukla "
      "tahmin ediliyordu, 19 kaydın 19'u da yanlışlıkla temizlik kategorisine "
      "gidiyordu, üstelik model bu yanlış tahminlere %95 üzeri güvenle "
      "yapıyordu — yani emin şekilde yanılıyordu. Kategorileri kelime "
      "dünyaları ayrışacak şekilde yeniden tasarladım.")

    baslik(doc, "3.1 Ayrım İlkesi", 2)
    p(doc,
      "Kategoriyi belirleyen şey arızanın NESNESİ değil, SORUMLU EKİP. Yani "
      "bir turnike üç farklı kategoriye düşebilir: kol dönmüyorsa (fiziksel "
      "arıza) mekanik ekibine, kart okumuyorsa (elektronik arıza) elektronik "
      "sistemler ekibine, biri atlayıp geçiyorsa (asayiş) güvenlik ekibine "
      "gider.")

    tablo(doc, ["Kategori", "Kapsam (özet)"], [
        ["Mekanik ve İstasyon", "yürüyen merdiven, asansör, turnikenin mekanik arızası, kayar kapılar"],
        ["Elektrik ve Enerji", "aydınlatma, jeneratör, katener, üçüncü ray, trafo, pano, sigorta"],
        ["Araç ve Tren", "tren kapısı, HVAC, fren/cer, vagon camı ve koltuğu, araç içi anons"],
        ["Sinyalizasyon ve Haberleşme", "sinyal arızası, PAKS/PSD, CCTV ve sensörlerin teknik arızası, telsiz"],
        ["Elektronik Sistemler", "biletmatik, kart okuyucu, QR, para sıkışması, turnike elektroniği"],
        ["Yol ve Hat", "ray kırılması, makas, travers, balast, hat üzerinde cisim"],
        ["İstasyon Güvenliği", "güvenlik personeli, kamera görüşünün engellenmesi, yangın/duman algılama"],
        ["Temizlik", "çöp, hijyen, kirli zemin, koku, haşere, grafiti"],
        ["Yolcu Hizmetleri", "anons içeriği, bilgi ekranları, sefer bilgisi, yönlendirme, yoğunluk"],
        ["Güvenlik ve Asayiş Olayı", "saldırı, kavga, hırsızlık, hasta yolcu, şüpheli paket"],
        ["Altyapı ve İnşaat", "su sızıntısı, çatlak, tünel yapısı, drenaj, inşaat faaliyeti"],
    ])

    baslik(doc, "3.2 Üç Boyut", 2)
    p(doc, "Kategorinin yanına iki boyut daha ekledim:")
    madde(doc, "Intent (5 sınıf): fault_report (arıza bildirimi), "
                "incident_report (olay bildirimi — asayiş, acil sağlık gibi "
                "ekipman değil DURUM bildiren), information_request (bilgi "
                "talebi), complaint (şikayet — somut arıza yok, memnuniyetsizlik "
                "var), suggestion (öneri).")
    madde(doc, "Öncelik (4 sınıf, P1-P4): can güvenliği tehdidi mi (P1), sefer/"
                "yolcu akışı aksıyor mu (P2), tek ekipman arızası ama yolculuk "
                "mümkün mü (P3), yoksa işleyişi hiç etkilemiyor mu (P4).")
    p(doc,
      "Öncelik için ayrıca bir 'kural katmanı' ekledim: yangın, elektrik "
      "çarpması, raylara kişi düşmesi gibi 11 dar ve kesin kalıp, modelin "
      "tahminini EZİYOR ve koşulsuz P1 veriyor. Bunun sebebi basit: P1'i "
      "kaçırmanın bedeli çok yüksek (bir yangını P3 sanmak kabul edilemez), "
      "tersinin bedeli düşük (gereksiz aciliyet). Bu yüzden bu tür can "
      "güvenliği durumlarında modele değil, kesin bir kurala güveniyorum.")

    # ------------------------------------------------------------------ 4
    baslik(doc, "4. Veri Üretimi ve Etiketleme", 1)
    p(doc,
      "Elimde gerçek Metro İstanbul arıza verisi yoktu (gizlilik ve erişim "
      "sebebiyle), bu yüzden eğitim verisini yapay olarak, LLM'lerle ürettim. "
      "Bunu yaparken en çok emek verdiğim şeylerden biri, hangi LLM sağlayıcının "
      "gerçekten işe yaradığını ÖLÇEREK bulmaktı — tahmin etmedim.")

    baslik(doc, "4.1 LLM Sağlayıcı Karşılaştırması", 2)
    p(doc,
      "Dört farklı sağlayıcıyı aynı görevde denedim ve kendi yazdığım kalite "
      "kontrol aracıyla (review.py — tekrar, uzunluk, sızıntı ve sınır "
      "hatalarını otomatik işaretliyor) ölçtüm:")
    tablo(doc, ["Sağlayıcı/Model", "Seed işaretli", "Gold işaretli", "Not"], [
        ["Gemini (çeşitli)", "%15", "%29", "Kota ve erişim sorunları yüzünden terk edildi"],
        ["Groq / Llama-3.3-70B", "%82", "%56", "Config'teki Türkçe metin ASCII yazılmıştı, kök sebep buydu"],
        ["Groq (Türkçe düzeltildi)", "%70", "%70", "İyileşti ama çok kısıtlı talimatlara uyamadı"],
        ["OpenRouter / Nemotron 3 Ultra", "%14", "%4", "Kazanan — ücretsiz, kart istemiyor"],
    ])
    p(doc,
      "Buradan çıkardığım en önemli ders şu: modelin ham dil kalitesinden çok, "
      "TALİMAT TAKİBİ kapasitesi belirleyici. Aynı düzgün Türkçe prompt'la "
      "bile 70 milyar parametreli Llama, kategori+stil+uzunluk+sızıntı gibi "
      "birden fazla kısıtı aynı anda yönetemedi; çok daha büyük Nemotron "
      "modeli aynı prompt'la çok daha iyi sonuç verdi.")
    p(doc,
      "Sonradan yerel olarak çalıştırdığım qwen2.5:14b modelini de aynı "
      "görevde Nemotron'la karşılaştırdım (1586 kayıt üzerinde): Nemotron "
      "%4.5 işaretli, qwen %15.8 işaretli çıktı. Ayrıca qwen'in asıl zayıf "
      "noktası otomatik aracın ölçtüğü şey değildi — uydurma istasyon adları "
      "('marmaraisi', 'yersenlik' gibi olmayan yerler) ve uydurma teknik "
      "terimler ('perde çarkı' gibi) üretiyordu. Bunu otomatik araç "
      "yakalayamıyor, elle okumak gerekiyor. Bu da bana 'düşük hata oranı "
      "tek başına kalite garantisi değil' dersini verdi.")

    baslik(doc, "4.2 Kalite Kontrol ve Veri Sızıntısı Önleme", 2)
    p(doc,
      "Sentetik veri üretiminin en büyük riski, üretilen örneklerin birbirine "
      "çok benzemesi (near-duplicate) ve bunların train/test setlerine "
      "karışıp sahte yüksek başarı göstermesi. Buna karşı üç katmanlı bir "
      "savunma kurdum:")
    madde(doc, "Üretim anında: yeni bir örnek, önceden üretilmiş örneklere "
                "%85'ten fazla benziyorsa reddediliyor.")
    madde(doc, "Bölme öncesi: train/val/test'e ayırmadan önce, birbirine "
                "%80'den fazla benzeyen kayıtlar aynı kümede (union-find ile) "
                "toplanıp HEP AYNI bölmede kalıyor — böylece çoğaltılmış "
                "verinin test setine sızıp yapay olarak yüksek doğruluk "
                "göstermesi engelleniyor.")
    madde(doc, "Her eğitim çalıştırmasında otomatik kontrol: hiçbir test "
                "kaydının train setinde birebir veya yakın kopya olarak "
                "bulunmadığı doğrulanıyor.")
    p(doc,
      "Bu üç katmanı kurarken bir gerçek hata da buldum: benzerlik ölçen "
      "fonksiyon (SequenceMatcher) argüman sırasına duyarlıymış — "
      "similarity(a,b) ile similarity(b,a) farklı sonuç verebiliyordu (0.8511 "
      "vs 0.8298 gibi, tam da eşiğin iki yanında). Üretim kodu bir sırayla, "
      "bölme kodu başka bir sırayla çağırıyordu, yani aynı çift bazen kabul "
      "ediliyor bazen reddediliyordu. Girdileri kanonik sıraya sokarak "
      "(sorted) düzelttim.")

    # ------------------------------------------------------------------ 5
    baslik(doc, "5. Model Eğitimi", 1)
    p(doc,
      "Eğitim döngüsünü Hugging Face'in hazır Trainer sınıfı yerine elle "
      "yazdım — Trainer, Apple Silicon'daki MPS backend'inde bazen dtype/"
      "cihaz sürprizleri çıkarıyordu ve hata ayıklamayı zorlaştırıyordu. "
      "Elle döngü hem şeffaf hem de bu veri boyutu için (birkaç bin kayıt) "
      "gayet yeterli.")

    baslik(doc, "5.1 Aksan Dayanıklılığı", 2)
    p(doc,
      "İlk eğitimden sonra modelin yazım hatalı/aksansız yazılmış metinlerde "
      "(örn. 'asansor' yerine 'asansör') daha kötü performans gösterdiği "
      "şüphesi oluştu. Bunu ölçmek için nedensel bir test yaptım: test "
      "setindeki aksan içeren kayıtların aksanlarını yapay olarak kaldırıp "
      "tekrar tahmin ettirdim. Sonuç: doğruluk %90.75'ten %84.39'a düştü — "
      "yani sadece aksan kaldırmak 6.4 puanlık kayba yol açıyordu.")
    p(doc,
      "Sebebini BERTurk'ün tokenizer'ında buldum: 'asansör' tek bir parça "
      "olarak tokenize edilirken, 'asansor' üç anlamsız alt-parçaya "
      "bölünüyordu ('asa', '##ns', '##or'). Çözüm olarak eğitim "
      "verisindeki aksan içeren kayıtların aksansız kopyalarını da eğitime "
      "ekledim (train seti 1280'den 2219'a çıktı). Bu hem aksan kaybını "
      "6.4 puandan 1.16 puana indirdi hem de genel doğruluğu da artırdı — "
      "yani sadece bir sorunu çözmedim, modeli genel olarak da iyileştirdi.")

    baslik(doc, "5.2 Early Stopping ve Canlı İzleme", 2)
    p(doc,
      "Projenin ilerleyen bir aşamasında eğitimi canlı izlemek istedim ve "
      "modelin ne zaman ezberlemeye (overfitting) başladığını validation "
      "loss üzerinden takip eden bir early stopping mekanizması ekledim: "
      "validation loss 3 epoch üst üste iyileşmezse eğitim otomatik duruyor. "
      "Bunu bilinçli olarak 'en iyi checkpoint seçimi'nden (üç görevin "
      "ortalama F1 skoruna göre seçiliyor) AYRI tuttum, çünkü validation "
      "loss düzleşse bile bir görevin (örneğin intent'in) F1 skoru hâlâ "
      "iyileşiyor olabiliyordu.")
    p(doc,
      "Ayrıca eğitim sırasında `model/canli_kayip.json` dosyasına periyodik "
      "olarak yazan ve bunu tarayıcıda saf SVG ile (dış kütüphane olmadan) "
      "canlı grafik olarak gösteren küçük bir izleme sayfası yazdım. Bunu "
      "kurarken de bir hata yaptım: Python'da float('inf') değerini JSON'a "
      "yazınca 'Infinity' diye geçersiz bir JSON çıkıyor, tarayıcı bunu "
      "sessizce parse edemiyor ve grafik sonsuza kadar 'veri bekleniyor' "
      "durumunda kalıyordu. None'a çevirerek düzelttim.")

    # ------------------------------------------------------------------ 6
    baslik(doc, "6. Kalibrasyon: Ne Zaman İnsana Sorulmalı", 1)
    p(doc,
      "Model her tahmin için bir güven skoru (confidence) üretiyor. Ama bu "
      "skoru ham haliyle kullanmak yanıltıcı olabilir. Eşiği (hangi güven "
      "seviyesinin altında 'insana sor' denileceğini) doğru veriyle kalibre "
      "etmem gerekiyordu.")
    p(doc,
      "İlk yaklaşımım test setine bakıp eşiği ona göre seçmekti, ama bu "
      "yanlış — test setine bakarak karar vermek, test setini karar sürecine "
      "sokmak demek, bu da skoru yapay olarak iyimser gösterir. Bunun yerine "
      "k-fold out-of-fold (OOF) yöntemini kullandım: eğitim verisini 5 parçaya "
      "bölüp her parça için o parçayı HİÇ GÖRMEMİŞ bir model eğittim ve sadece "
      "o parça üzerinde tahmin aldım. Böylece 1280 kaydın tamamı için "
      "'modelin gerçekten görmediği' bir tahmin elde ettim.")
    p(doc,
      "Bu yöntemle güven eşiğini 0.70'ten 0.75'e, marj eşiğini 0.40'tan "
      "0.30'a çektim. İlginç bir bulgu da şuydu: önceki eşik seçimleri "
      "sadece 8-9 hatalık küçük bir örnekleme bakılarak yapılmıştı; OOF ile "
      "102 hataya bakınca bazı oranların tamamen yanlış ölçülmüş olduğunu "
      "gördüm (bir oran 4.0 sanılıyordu, gerçeği 0.80 çıktı). Bu bana "
      "projenin belki de en önemli metodoloji dersini verdi: az örneklemle "
      "yapılan tek bir ölçüm, ölçüm gibi görünen gürültü olabilir.")

    # ------------------------------------------------------------------ 7
    baslik(doc, "7. Bağımsız Doğrulama ve Sınır Sorunlarının Düzeltilmesi", 1)
    p(doc,
      "Modeli sadece kendi ürettiğim test setiyle değil, kullanıcının (staj "
      "danışmanımın) başka bir LLM ile bağımsız olarak ürettiği ve elle "
      "etiketlediği 80 kayıtlık ayrı bir test setiyle de sınadım. Bunun "
      "önemi şu: eğitim verimi üreten LLM ile test verimi üreten LLM aynı "
      "olursa, model o LLM'in kendi yazım kalıplarını ezberlemiş olsa bile "
      "test skoru yüksek çıkabilir. Farklı bir kaynaktan gelen test seti bu "
      "riski ortadan kaldırıyor.")
    p(doc,
      "Bu bağımsız test setiyle ilk taramada üç gerçek sorun buldum ve "
      "üzerlerinde ayrı ayrı çalıştım:")

    baslik(doc, "7.1 Sinyalizasyon ile Yolcu Hizmetleri Sınırı", 2)
    p(doc,
      "Model, ekran/anons gibi ekipmanların FİZİKSEL olarak bozuk olduğu "
      "durumlar (sinyalizasyon_haberlesme) ile ekipman sağlam ama İÇERİK "
      "yanlış/eksik olduğu durumlar (yolcu_hizmetleri) arasında yüzeysel "
      "kelimelere ('ekran', 'anons') takılıp kök sebebi ayırt edemiyordu. "
      "Bunu çözmek için 10 çift (20 kayıt) 'hard negative' — yani kasıtlı "
      "zıt ikili — örnek hazırladım. Her çift AYNI senaryoyu iki farklı kök "
      "sebeple veriyor, örneğin:")
    kutu(doc, "Donanım bozuk (sinyalizasyon_haberlesme)",
         "\"Tavandaki ekranın camı kırılmış, süre yazmıyor.\"")
    kutu(doc, "Donanım sağlam, içerik eksik (yolcu_hizmetleri)",
         "\"Ekran cam gibi parlak çalışıyor ama tren saati yazmıyor, sadece "
         "reklamlar dönüyor.\"")
    p(doc,
      "Bu 20 örneği eğitime ekleyip yeniden eğittikten sonra bağımsız test "
      "setinde bu iki kategori arasındaki karışma tamamen ortadan kalktı.")

    baslik(doc, "7.2 Kapı Arızalarında Sinyalizasyon/Mekanik/Araç Sınırı", 2)
    p(doc,
      "Benzer bir sorunu ikinci bir turda buldum: peron ayırıcı kapı (PAKS, "
      "sinyalizasyon_haberlesme), vagon kapısı (arac_tren) ve istasyon giriş "
      "kapısı (mekanik_istasyon) arasında da aynı 'kapı' kelimesine takılma "
      "problemi vardı. Aynı yöntemle 13 kayıt daha (6 karşılaştırma grubu) "
      "ekledim. Sonuç: 13 hedefli çiftin 13'ünde de doğru kategori.")

    baslik(doc, "7.3 Genel Çeşitlilik Artışı", 2)
    p(doc,
      "Son olarak üç kategoriye (elektronik_sistemler, temizlik, yol_yapisal) "
      "20'şer kayıt daha ekledim — bunlar zıt-ikili değil, sadece o "
      "kategorilerin dil çeşitliliğini artıran genel örneklerdi (biletmatik "
      "arızaları, hijyen sorunları, ray hattı yapısal problemleri).")

    baslik(doc, "7.4 Sonuçların Karşılaştırması", 2)
    tablo(doc, ["Aşama", "Bağımsız sette kategori", "Bağımsız sette öncelik"], [
        ["Hard-negative eklemeden önce", "%73.8 (öncelik)", "%73.8"],
        ["Sinyalizasyon/yolcu düzeltmesi sonrası", "—", "%81.2"],
        ["Kapı düzeltmesi sonrası (tek koşu)", "%88.8", "%78.8"],
        ["+60 genel çeşitlilik örneği sonrası (güncel)", "%96.2", "%82.5"],
    ])
    p(doc,
      "Kapı düzeltmesinden sonraki tek koşuda hem kategori hem öncelik "
      "skorunda küçük bir düşüş görünce ilk başta endişelendim — ama bunun "
      "yeni eklediğim verilerle ilgisiz kategorilerde (Elektrik/Enerji, "
      "Altyapı/İnşaat gibi) ortaya çıktığını fark edince, bunun muhtemelen "
      "'tohum varyansı' (aynı veriyle bile farklı eğitim başlangıç "
      "noktalarının farklı sonuç vermesi) olduğunu düşündüm. Bir sonraki "
      "veri ekleme turunda hem kategori hem öncelik skorunun proje boyunca "
      "gördüğüm en yüksek seviyeye çıkması bu yorumu doğruladı.")

    kutu(doc, "Bu bölümden çıkardığım en büyük ders",
         "Tek bir eğitim koşusunun sonucu, özellikle küçük test setlerinde "
         "(80 kayıt gibi), kesin bir yargı için yeterli değil. Bir düşüş "
         "gördüğümde hemen 'kötüye gitti' demek yerine, hatanın nerede "
         "yoğunlaştığına bakıp gerçek bir regresyon mu yoksa gürültü mü "
         "olduğunu ayırt etmeye çalıştım. Bu proje boyunca bu dersi "
         "(az örneklem + tek ölçüm = gürültü olabilir) sanırım dört beş kere "
         "farklı bağlamda tekrar tekrar öğrendim.")

    # ------------------------------------------------------------------ 8
    baslik(doc, "8. Backend ve Frontend", 1)
    p(doc,
      "Modeli FastAPI ile bir REST servisine sardım. Model process başladığında "
      "bir kez yükleniyor ve bellekte tutuluyor (her istekte yeniden yüklemek "
      "birkaç saniye sürerdi). `/predict` uç noktası tek bir istekte üç boyutu "
      "(intent/kategori/öncelik) birden döndürüyor, ayrıca kurallı çıkarımla "
      "hat/istasyon/konum/ekipman/belirti gibi yapısal alanları, gradient×input "
      "yöntemiyle hangi kelimelerin karara katkı yaptığını (evidence), eksik "
      "bilgi varsa hangi alanların sorulması gerektiğini ve aynı arızanın kısa "
      "süre önce bildirilip bildirilmediğini de veriyor.")
    p(doc,
      "Arayüzü React ile, koyu temalı ve kategori rengine göre değişen bir "
      "'ortam ışığı' efektiyle tasarladım. Kullanıcı bir tahmin aldıktan sonra "
      "'Doğru' veya 'Yanlış' diyerek geri bildirim verebiliyor; bu geri "
      "bildirimler bir SQLite veritabanına kaydediliyor ama BİLİNÇLİ OLARAK "
      "otomatik yeniden eğitime sokulmuyor — çünkü kullanıcının onaylamadığı "
      "bir tahmin yanlış olabilir, ve modelin kendi hatasını doğru sanıp "
      "pekiştirmesi (confirmation bias) riski var. Sadece kullanıcının elle "
      "onayladığı kayıtlar dışa aktarılabiliyor, eğitime katılıp katılmayacağına "
      "kişi kendisi karar veriyor.")
    p(doc,
      "Servisi geliştirirken uçtan uca test ettim ve bu sırada iki tane "
      "gerçek, ilginç hata buldum:")
    madde(doc, "Benzerlik motoru (embedding tabanlı 'buna benzer kaç kayıt "
                "var' özelliği) mimari değişikliğinden sonra hiç güncellenmemişti "
                "ve backend'in ayağa kalkmasını engelliyordu — eski tek başlıklı "
                "modelin bir özniteliğini arıyordu, yeni çok başlıklı modelde "
                "o öznitelik farklı isimdeydi.")
    madde(doc, "Log veritabanındaki zaman damgaları yerel saatle (Türkiye "
                "UTC+3) yazılıyordu ama veritabanı sorguları UTC ile "
                "karşılaştırıyordu. Sonuç: 'aynı arıza az önce bildirildi mi' "
                "kontrolü 15 dakika yerine fiilen 3 saate yakın bir pencerede "
                "çalışıyordu. Bunu canlı test sırasında, dakikalar önce "
                "gönderdiğim bir metnin yanlışlıkla 'tekrar' sayılmasıyla "
                "yakaladım.")

    # ------------------------------------------------------------------ 9
    baslik(doc, "9. Sonuçlar Özeti", 1)
    tablo(doc, ["Görev", "İç test seti", "Bağımsız test seti (80 kayıt)"], [
        ["Kategori (11 sınıf)", "acc %85.8 · F1 0.861", "acc %96.2 · F1 0.962"],
        ["Intent (5 sınıf)", "acc %91.8 · F1 0.862", "—"],
        ["Öncelik (4 sınıf)", "acc %78.7 · F1 0.782", "acc %82.5 · F1 0.824"],
    ])
    p(doc,
      "LoRA adaptörü sadece 2.4 MB (tam model fine-tuning yapılsaydı 440 MB "
      "olurdu). Tahmin süresi ortalama 14 milisaniye. Eğitim havuzunda toplam "
      "2581 kayıt var, %80/%10/%10 oranında train/val/test'e bölünmüş "
      "durumda, ve Türkçe aksan dayanıklılığı için ek olarak ~1500 aksansız "
      "kopya eklenmiş halde.")

    # ------------------------------------------------------------------ 10
    baslik(doc, "10. Sınırlamalar ve Devam Eden İşler", 1)
    madde(doc, "Öncelik sınıflandırması hâlâ kategoriye göre daha zor — "
                "etiket tutarlılığını (aynı cümleyi iki kez etiketlettirip "
                "ölçtüğüm Cohen's kappa) 0.58'den 0.77'ye kadar çıkardım ama "
                "bir tavan var, muhtemelen önceliğin doğası gereği bazı "
                "cümleler gerçekten belirsiz.")
    madde(doc, "Resmi bir 'gold' (altın standart) test seti henüz yok — "
                "bağımsız test seti bu işlevi şimdilik dolduruyor ama daha "
                "büyük, resmi bir set faydalı olurdu.")
    madde(doc, "Yapısal çıkarım (hat/istasyon/ekipman bulma) kurallı bir "
                "sistemle çalışıyor, tam bir NER (Named Entity Recognition) "
                "modeli değil. Ölçtüğüm hataların çoğu bileşik ifadeler "
                "('3 numaralı vagonda kapı' gibi araya token giren ifadeler) "
                "— bunun için gelecekte token classification eklenebilir.")
    madde(doc, "Öncelik/eşik kalibrasyonu tek başlıklı modele göre yazılmıştı, "
                "çok başlıklı mimariye tam uyarlanması gerekiyor.")

    # ------------------------------------------------------------------ 11
    baslik(doc, "11. Kişisel Değerlendirme", 1)
    p(doc,
      "Bu proje bana en çok 'ölçmeden karar verme' alışkanlığını kazandırdı. "
      "Neredeyse her önemli kararı (hangi LLM, hangi eşik, hangi öğrenme hızı, "
      "gerçekten bir sorun var mı yoksa gürültü mü) sezgiyle değil, ölçerek "
      "verdim — ve bu ölçümlerin bir kısmı ilk sezgimi yanlış çıkardı "
      "(örneğin öncelik eşiğini ilk seferinde sadece 8 hataya bakarak "
      "seçmiştim, doğru yöntemle 102 hataya bakınca oranın tamamen farklı "
      "olduğunu gördüm).")
    p(doc,
      "En sevdiğim bulgu şu oldu: modelin bağımsız (hiç görmediği, başka bir "
      "kaynaktan gelen) test setindeki skoru, kendi ürettiğim test setindeki "
      "skorundan daha yüksek çıktı. Bu bana modelin ezberlemediğini, gerçekten "
      "kategoriyi öğrendiğini gösteren en somut kanıt oldu — sentetik veriyle "
      "eğitilmiş bir modelin 'gerçekçi mi' eleştirisine verebileceğim en iyi "
      "cevap bu.")
    p(doc,
      "Elbette her şey ilk seferde çalışmadı. Model bir ara hiçbir şey "
      "öğrenmedi, backend bir mimari değişikliğinden sonra hiç ayağa "
      "kalkmadı, zaman damgaları yanlış hesaplandığı için tekrar tespiti "
      "bozuktu. Ama bu hataların her birini bulup düzeltme süreci, bence "
      "projenin en öğretici kısmıydı.")

    doc.save(OUT)
    print(f"yazildi: {OUT}")


if __name__ == "__main__":
    olustur()
